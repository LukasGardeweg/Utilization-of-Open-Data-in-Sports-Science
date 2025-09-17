import docx
import random, numpy as np
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import re
from collections import Counter
import spacy
from spacy.lang.en.stop_words import STOP_WORDS

# Set random seed for reproducibility
random.seed(42)
np.random.seed(42)

# Load SpaCy model
nlp = spacy.load("en_core_web_md")

# Define keywords of interest
keywords = {
    "performance evaluation","performance", "evaluation", "quality", "scouting", "talents", "talent", "tactic analysis", "tactic", "analysis", "discovery",
    "match analysis", "match preparation", "match review", "review", "match scouting",
    "player evaluation", "player scouting", "player performance", "player development", "development", "team performance",
    "team scouting", "team tactics", "tactics", "comparison", "validation", "Success", "outcome", "prediction", "prediction of success", "success prediction", "Predictability", "predictive modeling",
    "complex systems", "complex", "global behavior", "behavior", "dynamics", "interaction", "interactions", "network", "networks",
    "passing network", "team strength", "dynamics of interaction", "team dynamics", "collaboration"
}

# Basis: SpaCy-Stopwords
all_stopwords = set(STOP_WORDS)

# own words to exclude
custom_stopwords = {"well", "more", "best"}
all_stopwords = all_stopwords.union(custom_stopwords)

# normalize text by lowercasing, removing punctuation, and extra spaces
def normalize_text(text):
    text = text.lower()                     # converting to lowercase
    text = re.sub(r"[-_/]", " ", text)      # replacing dashes/underscores/slashes with spaces
    text = re.sub(r"[^\w\s]", "", text)     # removing punctuation
    text = re.sub(r"\s+", " ", text)        # reducing multiple spaces to a single one
    return text.strip()

# Extract keywords/phrases from a Word document
def extract_phrases_from_docx(filepath, keywords):
    doc = docx.Document(filepath)
    found_phrases = []
    keywords_lower = {k.lower() for k in keywords}

    for para in doc.paragraphs:
        normalized_para = normalize_text(para.text)
        for phrase in keywords_lower:
            if phrase in normalized_para:
                # Leerzeichen in Phrase durch Unterstrich ersetzen,
                # damit WordCloud diese als ein Wort erkennt
                phrase_underscored = phrase.replace(" ", "_")
                found_phrases.append(phrase_underscored)
    return found_phrases

# Extract all normalized words from a Word document
def get_all_words_from_docx(filepath):
    doc = docx.Document(filepath)
    words = []
    for para in doc.paragraphs:
        normalized_para = normalize_text(para.text)
        for w in normalized_para.split():
            if w not in all_stopwords:  
                words.append(w)
    return words

# Find similar words to the keywords
def find_similar_words(words_list, keywords, similarity_threshold=1):
    keywords_docs = [nlp(k) for k in keywords]
    similar_words = []
    for word in words_list:  
        token_doc = nlp(word)
        for kw_doc in keywords_docs:
            if token_doc.similarity(kw_doc) >= similarity_threshold:
                similar_words.append(word)   
                break
    return similar_words


# path to word file (loading)
docx_path = r"C:\Users\Acer\Documents\Dokumente\Studium Bachelor Sport und Leistung\Bachelorarbeit\Wordcloud_abstract.docx"

found_words = extract_phrases_from_docx(docx_path, keywords)

# Find similar words to the keywords
all_words = get_all_words_from_docx(docx_path)
similar_words = find_similar_words(all_words, keywords)

# Combine found keywords and similar words
all_words_wordcloud = list(found_words) + list(similar_words)
text = " ".join(all_words_wordcloud)

if text:
    wc = WordCloud(width=800, height=400, background_color="white", max_words=50).generate(text)
    plt.figure(figsize=(10, 5))
    plt.imshow(wc, interpolation="bilinear")
    plt.axis("off")
    plt.show()
else:
    print("No keywords found.")

def count_word_frequencies(words):
    return Counter(words)

# Funktion zur Lemmatisierung (Singular/Plural etc. angleichen)
def lemmatize_word(word):
    doc = nlp(word)
    return doc[0].lemma_ if doc else word

# Frequenzzählung mit Lemmatisierung
def count_word_frequencies_lemmatized(words):
    lemmatized_words = [lemmatize_word(w) for w in words]
    return Counter(lemmatized_words)

# Frequenzen berechnen
frequencies = count_word_frequencies_lemmatized(all_words_wordcloud)

# Nur Wörter in der WordCloud berücksichtigen (ebenfalls lemmatisiert)
wc_words = set([lemmatize_word(w) for w in wc.words_.keys()])

print("Frequency of words in the WordCloud (lemmatized & sorted):")
for word, count in sorted(frequencies.items(), key=lambda x: x[1], reverse=True):
    if word in wc_words:
        print(f"{word}: {count}")

